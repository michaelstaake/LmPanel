from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

MAX_ATTACHMENT_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 120_000
TRUNCATION_SUFFIX = "\n\n[Attachment text truncated to fit the extraction limit.]"

TEXT_ATTACHMENT_SUFFIXES = {
    ".conf",
    ".css",
    ".csv",
    ".env",
    ".html",
    ".ini",
    ".js",
    ".json",
    ".log",
    ".md",
    ".py",
    ".sh",
    ".toml",
    ".ts",
    ".tsx",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
}

DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

XLSX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

ODT_MIME_TYPES = {
    "application/vnd.oasis.opendocument.text",
}

ODS_MIME_TYPES = {
    "application/vnd.oasis.opendocument.spreadsheet",
}

PDF_MIME_TYPES = {"application/pdf"}

ODF_NAMESPACES = {
    "table": "urn:oasis:names:tc:opendocument:xmlns:table:1.0",
    "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
}


def extract_attachment_text(filename: str, content_type: str | None, payload: bytes) -> dict:
    size = len(payload)
    normalized_content_type = (content_type or "").strip() or None
    suffix = Path(filename).suffix.lower()

    if size == 0:
        return _error_result(filename, normalized_content_type, size, "The file is empty.")

    if size > MAX_ATTACHMENT_BYTES:
        return _error_result(
            filename,
            normalized_content_type,
            size,
            f"File exceeds the {MAX_ATTACHMENT_BYTES // (1024 * 1024)} MB extraction limit.",
        )

    try:
        if _is_text_like(suffix, normalized_content_type):
            text = _decode_text_payload(payload)
            return _success_result(filename, normalized_content_type, size, text, extractor="text")

        if suffix == ".pdf" or normalized_content_type in PDF_MIME_TYPES:
            text = _extract_pdf_text(payload)
            if not text.strip():
                return _error_result(
                    filename,
                    normalized_content_type,
                    size,
                    "No extractable text was found in the PDF. Scanned PDFs need OCR, which is not enabled yet.",
                )
            return _success_result(filename, normalized_content_type, size, text, extractor="pdf")

        if suffix == ".docx" or normalized_content_type in DOCX_MIME_TYPES:
            text = _extract_docx_text(payload)
            if not text.strip():
                return _error_result(filename, normalized_content_type, size, "No extractable text was found in the DOCX file.")
            return _success_result(filename, normalized_content_type, size, text, extractor="docx")

        if suffix == ".xlsx" or normalized_content_type in XLSX_MIME_TYPES:
            text = _extract_xlsx_text(payload)
            if not text.strip():
                return _error_result(filename, normalized_content_type, size, "No extractable text was found in the XLSX file.")
            return _success_result(filename, normalized_content_type, size, text, extractor="xlsx")

        if suffix == ".odt" or normalized_content_type in ODT_MIME_TYPES:
            text = _extract_odt_text(payload)
            if not text.strip():
                return _error_result(filename, normalized_content_type, size, "No extractable text was found in the ODT file.")
            return _success_result(filename, normalized_content_type, size, text, extractor="odt")

        if suffix == ".ods" or normalized_content_type in ODS_MIME_TYPES:
            text = _extract_ods_text(payload)
            if not text.strip():
                return _error_result(filename, normalized_content_type, size, "No extractable text was found in the ODS file.")
            return _success_result(filename, normalized_content_type, size, text, extractor="ods")
    except Exception as exc:
        return _error_result(filename, normalized_content_type, size, str(exc) or "Text extraction failed.")

    return {
        "name": filename,
        "contentType": normalized_content_type,
        "size": size,
        "status": "unsupported",
        "content": None,
        "detail": "This file type is not supported for text extraction yet.",
        "truncated": False,
        "extractor": None,
    }


def _is_text_like(suffix: str, content_type: str | None) -> bool:
    return suffix in TEXT_ATTACHMENT_SUFFIXES or (content_type is not None and content_type.startswith("text/"))


def _decode_text_payload(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue

    return payload.decode("utf-8", errors="replace")


def _extract_pdf_text(payload: bytes) -> str:
    reader = PdfReader(BytesIO(payload))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(page for page in pages if page.strip())


def _extract_docx_text(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    segments = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                segments.append(row_text)

    return "\n".join(segments)


def _extract_xlsx_text(payload: bytes) -> str:
    workbook = load_workbook(filename=BytesIO(payload), read_only=True, data_only=True)
    sections: list[str] = []

    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                rows.append(" | ".join(values))

        if rows:
            sections.append(f"[Sheet: {sheet.title}]\n" + "\n".join(rows))

    return "\n\n".join(sections)


def _extract_odt_text(payload: bytes) -> str:
    root = _load_odf_content_xml(payload)
    paragraphs: list[str] = []

    for element in root.findall(".//text:p", ODF_NAMESPACES):
        text = _flatten_xml_text(element)
        if text:
            paragraphs.append(text)

    for element in root.findall(".//text:h", ODF_NAMESPACES):
        text = _flatten_xml_text(element)
        if text:
            paragraphs.append(text)

    return "\n".join(paragraphs)


def _extract_ods_text(payload: bytes) -> str:
    root = _load_odf_content_xml(payload)
    sections: list[str] = []

    for table in root.findall(".//table:table", ODF_NAMESPACES):
        sheet_name = table.attrib.get(f"{{{ODF_NAMESPACES['table']}}}name", "Sheet")
        rows: list[str] = []

        for row in table.findall("table:table-row", ODF_NAMESPACES):
            cells: list[str] = []
            for cell in row.findall("table:table-cell", ODF_NAMESPACES):
                repeated_columns = _parse_positive_int(cell.attrib.get(f"{{{ODF_NAMESPACES['table']}}}number-columns-repeated"))
                text = _flatten_xml_text(cell)
                if text:
                    cells.extend([text] * repeated_columns)

            if cells:
                rows.append(" | ".join(cells))

        if rows:
            sections.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))

    return "\n\n".join(sections)


def _load_odf_content_xml(payload: bytes) -> ET.Element:
    with ZipFile(BytesIO(payload)) as archive:
        content_xml = archive.read("content.xml")
    return ET.fromstring(content_xml)


def _flatten_xml_text(element: ET.Element) -> str:
    text = "".join(segment.strip() for segment in element.itertext())
    return " ".join(text.split())


def _parse_positive_int(value: str | None) -> int:
    if not value:
        return 1

    try:
        parsed = int(value)
    except ValueError:
        return 1

    return parsed if parsed > 0 else 1


def _success_result(filename: str, content_type: str | None, size: int, text: str, extractor: str) -> dict:
    normalized_text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "").strip()
    truncated = False

    if len(normalized_text) > MAX_EXTRACTED_CHARACTERS:
        normalized_text = normalized_text[:MAX_EXTRACTED_CHARACTERS].rstrip() + TRUNCATION_SUFFIX
        truncated = True

    return {
        "name": filename,
        "contentType": content_type,
        "size": size,
        "status": "ok",
        "content": normalized_text,
        "detail": None,
        "truncated": truncated,
        "extractor": extractor,
    }


def _error_result(filename: str, content_type: str | None, size: int, detail: str) -> dict:
    return {
        "name": filename,
        "contentType": content_type,
        "size": size,
        "status": "error",
        "content": None,
        "detail": detail,
        "truncated": False,
        "extractor": None,
    }